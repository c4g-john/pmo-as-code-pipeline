#!/usr/bin/env python3
"""Extract plain text from a source document for doc-to-pmo conversion.

Usage:
    python tools/extract.py path/to/source.docx

Supports .docx, .pdf, .md, .txt. This is the deterministic first step of the
conversion front-door: it turns an arbitrary source file into plain text that
the doc-to-pmo skill then maps into a standard template. It does not interpret
or reshape the content — that is the skill's job.
"""
from __future__ import annotations

import sys
from pathlib import Path


def extract(path: str | Path) -> str:
    p = Path(path)
    if not p.is_file():
        raise SystemExit(f"extract: no such file: {p}")
    ext = p.suffix.lower()

    if ext in {".md", ".txt"}:
        return p.read_text(encoding="utf-8")

    if ext == ".docx":
        try:
            import docx  # python-docx
        except ImportError:
            raise SystemExit("extract: install support with  pip install '.[convert]'")
        document = docx.Document(str(p))
        blocks: list[str] = [para.text for para in document.paragraphs]
        # include table cell text, which charters often use for milestones/risks
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    blocks.append(" | ".join(cells))
        return "\n".join(blocks)

    if ext == ".pdf":
        try:
            from pypdf import PdfReader
        except ImportError:
            raise SystemExit("extract: install support with  pip install '.[convert]'")
        reader = PdfReader(str(p))
        return "\n".join((page.extract_text() or "") for page in reader.pages)

    raise SystemExit(f"extract: unsupported source type '{ext}' "
                     f"(supported: .docx, .pdf, .md, .txt)")


if __name__ == "__main__":
    if len(sys.argv) != 2:
        raise SystemExit("usage: python tools/extract.py <source-file>")
    sys.stdout.write(extract(sys.argv[1]))
